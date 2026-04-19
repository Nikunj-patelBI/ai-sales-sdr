"""
Generate Word document: AI-Native Sales Pipeline — Technical Architecture & Learning Blueprint.
This doc designs the sales pipeline using MCP, Agents, RAG, Vector DBs so the founders
learn cutting-edge AI by building, and can showcase it as a portfolio case study.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Colors
INDIGO = RGBColor(0x4F, 0x46, 0xE5)
DARK_INDIGO = RGBColor(0x31, 0x2E, 0x81)
TEAL = RGBColor(0x0D, 0x94, 0x88)
PURPLE = RGBColor(0x7C, 0x3A, 0xED)
DARK = RGBColor(0x1E, 0x1E, 0x2E)
GRAY = RGBColor(0x64, 0x64, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ORANGE = RGBColor(0xEA, 0x58, 0x0C)

# ── Helpers ──
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_table_borders(table, color="E2E8F0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def add_styled_table(headers, rows, col_widths=None, header_bg="F1F5F9"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_INDIGO
        set_cell_shading(cell, header_bg)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK
            if r_idx % 2 == 1:
                set_cell_shading(cell, "FAFAFA")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_heading_styled(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(22)
        run.font.color.rgb = DARK_INDIGO
        run.bold = True
        p.space_before = Pt(24)
        p.space_after = Pt(8)
        border_p = doc.add_paragraph()
        border_p.space_before = Pt(0)
        border_p.space_after = Pt(12)
        pPr = border_p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="7C3AED"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = PURPLE
        run.bold = True
        p.space_before = Pt(18)
        p.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = TEAL
        run.bold = True
        p.space_before = Pt(12)
        p.space_after = Pt(4)
    return p

def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    p.paragraph_format.line_spacing = Pt(16)
    p.space_after = Pt(6)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = DARK
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK
    else:
        p.text = ""
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(2)
    return p

def add_callout_box(text, title=None, border_color="7C3AED", bg_color="F5F3FF"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg_color)
    if title:
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = DARK_INDIGO
        p2 = cell.add_paragraph()
        r2 = p2.add_run(text)
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK
    else:
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = DARK
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'  <w:left w:val="single" w:sz="12" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    doc.add_paragraph()

def add_numbered(items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = PURPLE
        r2 = p.add_run(item)
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(2)

def add_code_block(text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "1E1E2E")
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0xA5, 0xD6, 0xFF)
    r.font.name = "Consolas"
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="3B3B5C"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="3B3B5C"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="3B3B5C"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="3B3B5C"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    doc.add_paragraph()

def add_page_break():
    doc.add_page_break()

def add_learning_badge(tech, what_you_learn):
    """Small callout showing what AI skill this section teaches."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, "DDD6FE")
    c1 = table.rows[0].cells[0]
    c2 = table.rows[0].cells[1]
    set_cell_shading(c1, "EDE9FE")
    set_cell_shading(c2, "F5F3FF")
    c1.width = Cm(3.5)
    c2.width = Cm(13)
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(f"LEARN: {tech}")
    r1.bold = True
    r1.font.size = Pt(8)
    r1.font.color.rgb = PURPLE
    p2 = c2.paragraphs[0]
    r2 = p2.add_run(what_you_learn)
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = DARK
    r2.italic = True
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ANALYTICSGEAR")
r.font.size = Pt(14)
r.font.color.rgb = PURPLE
r.bold = True
r.font.letter_spacing = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AI-Native Sales Pipeline")
r.font.size = Pt(36)
r.font.color.rgb = DARK_INDIGO
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Built with MCP, AI Agents, RAG & Vector Databases")
r.font.size = Pt(18)
r.font.color.rgb = PURPLE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("━" * 40)
r.font.color.rgb = PURPLE
r.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "A technical architecture that serves dual purpose:\n"
    "Learn cutting-edge AI by building  |  Showcase as a real-world AI case study"
)
r.font.size = Pt(11)
r.font.color.rgb = GRAY
r.italic = True

for _ in range(4):
    doc.add_paragraph()

meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate([
    ("Document Version", "1.0"),
    ("Date", datetime.date.today().strftime("%B %d, %Y")),
    ("Purpose", "Technical Architecture + Learning Path + Portfolio Case Study"),
    ("Classification", "Internal / Confidential"),
]):
    c1 = meta_table.rows[i].cells[0]
    c2 = meta_table.rows[i].cells[1]
    c1.text = ""
    c2.text = ""
    r1 = c1.paragraphs[0].add_run(label)
    r1.font.size = Pt(9)
    r1.font.color.rgb = GRAY
    r1.bold = True
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = c2.paragraphs[0].add_run(value)
    r2.font.size = Pt(9)
    r2.font.color.rgb = DARK

add_page_break()

# ════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════

add_heading_styled("Table of Contents", 1)

toc_items = [
    "Why Build This Way — The Dual-Purpose Strategy",
    "AI Concepts Primer — MCP, Agents, RAG, Vector DBs",
    "System Architecture — The AI-Native Pipeline",
    "Component 1 — Vector Knowledge Base (RAG + Vector DB)",
    "Component 2 — AI Agent Orchestrator (Multi-Agent System)",
    "Component 3 — MCP Integration Layer (Tool Use)",
    "Component 4 — Intelligent Conversation Memory",
    "Component 5 — Self-Improving Feedback Loop",
    "Tech Stack & Infrastructure",
    "Build-and-Learn Roadmap (12 Weeks)",
    "Turning This Into a Case Study",
    "Portfolio Presentation Guide",
    "What You'll Be Able to Sell After Building This",
]

for i, item in enumerate(toc_items, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"  {i:02d}   ")
    r.font.size = Pt(10)
    r.font.color.rgb = PURPLE
    r.bold = True
    r2 = p.add_run(item)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(4)

add_page_break()

# ════════════════════════════════════════════════════════════════════
#  1. WHY BUILD THIS WAY
# ════════════════════════════════════════════════════════════════════

add_heading_styled("1. Why Build This Way — The Dual-Purpose Strategy", 1)

add_body(
    "Most founders learn AI by taking courses or building toy projects that never see production. "
    "This approach is different. You build a real, production system that generates revenue for "
    "AnalyticsGear while learning every major AI pattern that clients are asking for."
)

add_callout_box(
    "BUILD IT — A real AI-native sales pipeline that runs daily and brings in clients\n"
    "LEARN IT — Hands-on experience with MCP, Agents, RAG, Vector DBs, and LLM orchestration\n"
    "SELL IT — A live case study that proves to prospects you can build what you're selling",
    title="The Three-in-One:"
)

add_heading_styled("The Problem with Learning AI in Isolation", 2)

add_body(
    "Reading about RAG architectures is not the same as building one that handles 10,000 "
    "company profiles and retrieves the right context in under 200ms. Watching an agent tutorial "
    "is not the same as debugging why your prospecting agent hallucinated a CEO's name. "
    "You need production pressure to truly learn."
)

add_heading_styled("What This Sales Pipeline Teaches You", 2)

add_styled_table(
    ["AI Technology", "Where It's Used in the Pipeline", "What You'll Master"],
    [
        ["RAG (Retrieval Augmented Generation)", "Retrieving company intel, past interactions, blog content for personalized outreach",
         "Chunking strategies, embedding models, retrieval quality, context window management"],
        ["Vector Databases", "Storing and searching lead profiles, company data, email templates, interaction history",
         "Embedding pipelines, similarity search, metadata filtering, index optimization"],
        ["AI Agents", "Autonomous prospecting agent, outreach agent, follow-up agent, analytics agent",
         "Agent architecture, tool use, planning, multi-agent coordination, error recovery"],
        ["MCP (Model Context Protocol)", "Connecting agents to LinkedIn, email, CRM, web scraping, and databases as tools",
         "MCP server development, tool schemas, resource management, protocol implementation"],
        ["LLM Orchestration", "Chaining lead scoring, email generation, reply handling, reporting",
         "Prompt engineering, chain design, fallback handling, cost optimization"],
        ["Evaluation & Feedback Loops", "Measuring which emails get replies, which scoring is accurate, which agents perform",
         "LLM evaluation, A/B testing, reinforcement from human feedback, continuous improvement"],
    ],
    col_widths=[4, 6, 6.5],
)

add_heading_styled("Why Clients Will Care", 2)

add_body(
    'When a prospect asks "Have you built a RAG system before?" you won\'t say "We did a proof of concept." '
    'You\'ll say: "We run one in production daily. It powers our sales pipeline — processes 500+ company '
    'profiles, generates personalized outreach, and has booked us 40+ meetings in 3 months. '
    'Here\'s how we built it." That\'s a fundamentally different conversation.'
)

add_page_break()

# ════════════════════════════════════════════════════════════════════
#  2. AI CONCEPTS PRIMER
# ════════════════════════════════════════════════════════════════════

add_heading_styled("2. AI Concepts Primer", 1)

add_body(
    "Before diving into the architecture, here's a concise primer on the four core technologies "
    "and how they relate to each other."
)

# RAG
add_heading_styled("2.1 RAG — Retrieval Augmented Generation", 2)

add_body(
    "RAG solves the problem of giving an LLM access to your private data without fine-tuning. "
    "Instead of training the model on your data, you retrieve relevant context at query time "
    "and inject it into the prompt."
)

add_body(
    "In our pipeline: When the AI writes an email to a prospect, it retrieves that company's "
    "profile, tech stack, recent news, past interaction history, and relevant AnalyticsGear "
    "case studies — all from a vector database — to write a deeply personalized message."
)

add_code_block(
    "HOW RAG WORKS:\n\n"
    "  [User Query]  -->  [Embed Query]  -->  [Search Vector DB]  -->  [Top-K Results]\n"
    "                                                                        |\n"
    "                                                                        v\n"
    "  [LLM Response]  <--  [LLM generates with context]  <--  [Prompt + Retrieved Docs]"
)

# Vector DB
add_heading_styled("2.2 Vector Databases", 2)

add_body(
    "A vector database stores data as high-dimensional numerical vectors (embeddings) rather than "
    "rows and columns. This enables semantic search — finding things by meaning rather than "
    "exact keyword match. When you search for 'companies struggling with data pipelines', it "
    "returns results about 'ETL bottlenecks' and 'broken Airflow DAGs' even though the words "
    "don't match."
)

add_styled_table(
    ["Vector DB", "Best For", "Pricing"],
    [
        ["Pinecone", "Managed, easiest to start, great for production", "Free tier, then $25/mo+"],
        ["Qdrant", "Self-hosted or cloud, rich filtering, fast", "Free (self-host) or cloud"],
        ["ChromaDB", "Local development, embedded, Python-native", "Free, open-source"],
        ["Weaviate", "Hybrid search (vector + keyword), multimodal", "Free (self-host) or cloud"],
        ["pgvector", "If you already use PostgreSQL", "Free extension"],
    ],
    col_widths=[3.5, 7, 6],
)

# Agents
add_heading_styled("2.3 AI Agents", 2)

add_body(
    "An AI agent is an LLM that can reason, plan, and take actions autonomously using tools. "
    "Unlike a simple chatbot that only generates text, an agent can: decide what information "
    "it needs, call APIs to get that information, process the results, decide the next step, "
    "and take action — all without human intervention."
)

add_body(
    "In our pipeline: A Prospecting Agent autonomously searches for leads, an Outreach Agent "
    "crafts and sends personalized messages, a Follow-Up Agent monitors responses and takes "
    "appropriate action, and an Analytics Agent generates weekly performance reports."
)

add_code_block(
    "AGENT LOOP:\n\n"
    "  [Goal: Find 25 qualified leads]           [Goal: Send personalized emails]\n"
    "          |                                            |\n"
    "          v                                            v\n"
    "  [Think: Search Apollo for CTOs]            [Think: Need company context]\n"
    "          |                                            |\n"
    "          v                                            v\n"
    "  [Act: Call Apollo API tool]                 [Act: Query Vector DB tool]\n"
    "          |                                            |\n"
    "          v                                            v\n"
    "  [Observe: Got 30 results]                   [Observe: Got company profile]\n"
    "          |                                            |\n"
    "          v                                            v\n"
    "  [Think: Need to score & filter]             [Think: Draft email with context]\n"
    "          |                                            |\n"
    "          v                                            v\n"
    "  [Act: Score with lead scoring tool]          [Act: Generate & send email]\n"
    "          |                                            |\n"
    "          v                                            v\n"
    "  [Done: 25 qualified leads saved]             [Done: 15 emails sent]"
)

# MCP
add_heading_styled("2.4 MCP — Model Context Protocol", 2)

add_body(
    "MCP is an open protocol (created by Anthropic) that standardizes how AI models connect to "
    "external tools and data sources. Think of it as a USB-C for AI — a universal interface "
    "that lets any AI model talk to any tool through a standard protocol."
)

add_body(
    "Instead of writing custom API integration code for each tool, you build MCP servers that "
    "expose tools in a standard format. Any MCP-compatible AI client can then discover and use "
    "those tools automatically."
)

add_body(
    "In our pipeline: We build MCP servers for LinkedIn, email (SendGrid), CRM (Google Sheets), "
    "Apollo.io, and our vector database. The AI agents connect to these tools via MCP, meaning "
    "adding a new tool (e.g., HubSpot) is just adding another MCP server — zero agent code changes."
)

add_code_block(
    "MCP ARCHITECTURE:\n\n"
    "  ┌─────────────┐      MCP Protocol       ┌──────────────────┐\n"
    "  |  AI Agent    | <---------------------> | MCP Server:      |\n"
    "  |  (Client)    |                         | LinkedIn         |\n"
    "  |              | <---------------------> | MCP Server:      |\n"
    "  |  Discovers   |                         | Email/SendGrid   |\n"
    "  |  tools auto- | <---------------------> | MCP Server:      |\n"
    "  |  matically   |                         | CRM/Sheets       |\n"
    "  |              | <---------------------> | MCP Server:      |\n"
    "  |              |                         | Vector DB        |\n"
    "  └─────────────┘                         └──────────────────┘"
)

add_page_break()

# ════════════════════════════════════════════════════════════════════
#  3. SYSTEM ARCHITECTURE
# ════════════════════════════════════════════════════════════════════

add_heading_styled("3. System Architecture — The AI-Native Pipeline", 1)

add_body(
    "This architecture replaces simple Python scripts with a multi-agent system backed by "
    "a vector knowledge base and connected to external services via MCP. The result is a "
    "pipeline that reasons, remembers, and improves over time."
)

# Architecture diagram as table
add_heading_styled("3.1 High-Level Architecture", 2)

arch = doc.add_table(rows=5, cols=3)
arch.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(arch, "DDD6FE")

# Row 0: Agent Layer
for i in range(3):
    set_cell_shading(arch.rows[0].cells[i], "EDE9FE")
labels_r0 = [
    ("PROSPECTING AGENT", "Finds & qualifies leads"),
    ("OUTREACH AGENT", "Crafts & sends messages"),
    ("FOLLOW-UP AGENT", "Monitors & responds"),
]
for i, (title, desc) in enumerate(labels_r0):
    p = arch.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title + "\n")
    r.bold = True; r.font.size = Pt(8); r.font.color.rgb = PURPLE
    r2 = p.add_run(desc)
    r2.font.size = Pt(7.5); r2.font.color.rgb = GRAY

# Row 1: Arrow
for i in range(3):
    p = arch.rows[1].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("▼ MCP Protocol ▼")
    r.font.size = Pt(7); r.font.color.rgb = PURPLE; r.bold = True

# Row 2: MCP Servers
set_cell_shading(arch.rows[2].cells[0], "F0FDFA")
set_cell_shading(arch.rows[2].cells[1], "F0FDFA")
set_cell_shading(arch.rows[2].cells[2], "F0FDFA")
mcp_labels = [
    ("MCP: LinkedIn + Apollo", "Search, connect, message"),
    ("MCP: Email + CRM", "SendGrid, Google Sheets"),
    ("MCP: Web Scraping", "Apify, job boards, news"),
]
for i, (title, desc) in enumerate(mcp_labels):
    p = arch.rows[2].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title + "\n")
    r.bold = True; r.font.size = Pt(8); r.font.color.rgb = TEAL
    r2 = p.add_run(desc)
    r2.font.size = Pt(7.5); r2.font.color.rgb = GRAY

# Row 3: Arrow
for i in range(3):
    p = arch.rows[3].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("▼")
    r.font.size = Pt(7); r.font.color.rgb = PURPLE; r.bold = True

# Row 4: Knowledge layer (merged conceptually)
set_cell_shading(arch.rows[4].cells[0], "EEF2FF")
set_cell_shading(arch.rows[4].cells[1], "EEF2FF")
set_cell_shading(arch.rows[4].cells[2], "EEF2FF")
kb_labels = [
    ("VECTOR DB", "Company profiles, embeddings"),
    ("RAG ENGINE", "Context retrieval for prompts"),
    ("MEMORY STORE", "Interaction history, learnings"),
]
for i, (title, desc) in enumerate(kb_labels):
    p = arch.rows[4].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title + "\n")
    r.bold = True; r.font.size = Pt(8); r.font.color.rgb = INDIGO
    r2 = p.add_run(desc)
    r2.font.size = Pt(7.5); r2.font.color.rgb = GRAY

doc.add_paragraph()

add_heading_styled("3.2 How the Components Connect", 2)

add_numbered([
    "The Orchestrator wakes up on schedule and activates agents based on the daily pipeline stage",
    "Each Agent plans its tasks, then calls tools via MCP to execute (search leads, send emails, etc.)",
    "Before any outreach, the agent queries the RAG engine which retrieves relevant context from the Vector DB",
    "All interactions are logged to the Memory Store, building a richer context over time",
    "The Analytics Agent reviews performance data and adjusts strategies (which emails work, which leads convert)",
    "The feedback loop makes the system smarter every week — scoring improves, messaging improves, targeting improves",
])

add_page_break()

# ════════════════════════════════════════════════════════════════════
#  4. COMPONENT 1 — VECTOR KNOWLEDGE BASE
# ════════════════════════════════════════════════════════════════════

add_heading_styled("4. Component 1 — Vector Knowledge Base", 1)

add_learning_badge("RAG + Vector DB", "Embeddings, chunking, similarity search, metadata filtering, retrieval quality tuning")

add_body(
    "The Vector Knowledge Base is the brain of the pipeline. It stores everything the agents "
    "need to know — company profiles, prospect data, AnalyticsGear's own content, past interactions, "
    "and successful outreach patterns — all as searchable vector embeddings."
)

add_heading_styled("4.1 What Gets Stored", 2)

add_styled_table(
    ["Collection", "Content", "Use Case", "Update Frequency"],
    [
        ["company_profiles", "Company name, industry, size, tech stack, funding, news, pain signals",
         "Personalizing outreach, scoring leads", "Daily (new leads)"],
        ["prospect_contacts", "Name, title, LinkedIn activity, connection status, interaction history",
         "Choosing who to contact and how", "Daily"],
        ["ag_content", "AnalyticsGear blog posts, case studies, service descriptions, engagement models",
         "RAG retrieval for email content — sharing relevant insights", "Weekly"],
        ["email_history", "Sent emails, open/click/reply data, which messaging worked",
         "Learning what works — the AI improves its emails over time", "Daily"],
        ["industry_knowledge", "Industry-specific pain points, regulations, trends, terminology",
         "Writing industry-aware outreach (banking compliance, retail analytics)", "Monthly"],
        ["successful_patterns", "Emails that got replies, LinkedIn messages that converted, winning sequences",
         "The AI learns from success and replicates patterns", "Weekly"],
    ],
    col_widths=[3.5, 5, 5, 3],
)

add_heading_styled("4.2 Embedding Pipeline", 2)

add_code_block(
    "EMBEDDING PIPELINE:\n\n"
    "  [Raw Data]  -->  [Chunking]  -->  [Embedding Model]  -->  [Vector DB]\n"
    "                       |                  |                       |\n"
    "               Split into            Generate                 Store with\n"
    "               meaningful            384/1536-dim              metadata\n"
    "               chunks                vectors                  for filtering\n\n"
    "  CHUNKING STRATEGY BY DATA TYPE:\n"
    "  - Company profiles:  One chunk per company (structured JSON -> text)\n"
    "  - Blog posts:        Paragraph-level chunks with overlap (200 tokens, 50 overlap)\n"
    "  - Email history:     One chunk per email with outcome metadata\n"
    "  - Industry knowledge: Section-level chunks with industry tag metadata"
)

add_heading_styled("4.3 Retrieval Flow (RAG in Action)", 2)

add_body("When the Outreach Agent needs to write an email to a prospect:")

add_numbered([
    "Agent receives task: 'Write email to Sarah Chen, CTO at DataFlow Inc'",
    "Agent queries Vector DB: 'DataFlow Inc company profile' → retrieves company data, tech stack, recent news",
    "Agent queries Vector DB: 'successful emails to CTOs in SaaS companies' → retrieves top-performing email patterns",
    "Agent queries Vector DB: 'AnalyticsGear content relevant to data pipeline challenges' → retrieves matching blog post",
    "Agent queries Vector DB: 'past interactions with DataFlow Inc' → retrieves any previous touchpoints",
    "All retrieved context is injected into the prompt alongside the task, and the LLM generates a deeply personalized email",
])

add_heading_styled("4.4 Implementation", 2)

add_code_block(
    'from qdrant_client import QdrantClient\n'
    'from qdrant_client.models import Distance, VectorParams, PointStruct\n'
    'import anthropic\n'
    'import voyageai  # or use OpenAI embeddings\n\n'
    '# Initialize\n'
    'qdrant = QdrantClient(path="./qdrant_data")  # Local for dev\n'
    'voyage = voyageai.Client()  # For embeddings\n'
    'claude = anthropic.Anthropic()  # For generation\n\n'
    '# Create collection\n'
    'qdrant.create_collection(\n'
    '    collection_name="company_profiles",\n'
    '    vectors_config=VectorParams(size=1024, distance=Distance.COSINE)\n'
    ')\n\n'
    '# Embed and store a company profile\n'
    'def store_company(company: dict):\n'
    '    text = f"{company[\'name\']} | {company[\'industry\']} | "\\\n'
    '           f"Tech: {company[\'tech_stack\']} | {company[\'description\']}"\n'
    '    embedding = voyage.embed([text], model="voyage-3").embeddings[0]\n'
    '    qdrant.upsert("company_profiles", points=[PointStruct(\n'
    '        id=company["id"], vector=embedding,\n'
    '        payload=company  # Store full data as metadata\n'
    '    )])\n\n'
    '# Retrieve relevant companies for outreach\n'
    'def find_similar_companies(query: str, limit=5):\n'
    '    query_vec = voyage.embed([query], model="voyage-3").embeddings[0]\n'
    '    results = qdrant.search("company_profiles",\n'
    '        query_vector=query_vec, limit=limit,\n'
    '        query_filter={"must": [{"key": "score", "range": {"gte": 50}}]}\n'
    '    )\n'
    '    return [hit.payload for hit in results]'
)

add_page_break()

# ════════════════════════════════════════════════════════════════════
#  5. COMPONENT 2 — AGENT ORCHESTRATOR
# ════════════════════════════════════════════════════════════════════

add_heading_styled("5. Component 2 — AI Agent Orchestrator", 1)

add_learning_badge("AI Agents", "Agent architecture, tool use, ReAct pattern, multi-agent coordination, planning, error recovery")

add_body(
    "The pipeline is powered by four specialized AI agents, each with a distinct role, "
    "tool access, and decision-making capability. An orchestrator coordinates their execution "
    "based on the daily schedule and pipeline state."
)

add_heading_styled("5.1 Agent Definitions", 2)

add_styled_table(
    ["Agent", "Role", "Tools (via MCP)", "Runs When"],
    [
        ["Prospecting Agent", "Find and qualify new leads matching ICP",
         "Apollo search, LinkedIn search, Job board scraper, Lead scorer, Vector DB write",
         "Daily 6-7 AM"],
        ["Outreach Agent", "Generate and send personalized outreach across channels",
         "Vector DB read (RAG), Email sender, LinkedIn messenger, Call brief generator",
         "Daily 8 AM"],
        ["Follow-Up Agent", "Monitor engagement signals and manage responses",
         "Email tracker, LinkedIn monitor, Reply classifier, Sequence manager, CRM updater",
         "Daily 9 AM + triggered"],
        ["Analytics Agent", "Analyze performance and recommend improvements",
         "CRM reader, Metrics calculator, Report generator, A/B test analyzer",
         "Friday afternoons"],
    ],
    col_widths=[3, 4, 5.5, 3],
)

add_heading_styled("5.2 Agent Architecture Pattern (ReAct)", 2)

add_body(
    "Each agent follows the ReAct (Reason + Act) pattern — a loop of thinking, acting, "
    "observing, and deciding the next step. This is the same pattern used by Claude, ChatGPT, "
    "and every major AI agent framework."
)

add_code_block(
    'class SalesAgent:\n'
    '    """Base agent using ReAct pattern with MCP tool access."""\n\n'
    '    def __init__(self, name, system_prompt, mcp_tools):\n'
    '        self.name = name\n'
    '        self.system_prompt = system_prompt\n'
    '        self.tools = mcp_tools  # Tools discovered via MCP\n'
    '        self.memory = []  # Conversation/action history\n\n'
    '    async def run(self, goal: str) -> dict:\n'
    '        """Execute the agent loop until goal is achieved."""\n'
    '        messages = [{"role": "user", "content": goal}]\n\n'
    '        while True:\n'
    '            # REASON: Ask Claude to think and decide next action\n'
    '            response = claude.messages.create(\n'
    '                model="claude-sonnet-4-6",\n'
    '                system=self.system_prompt,\n'
    '                messages=messages,\n'
    '                tools=self.tools,  # MCP-provided tool schemas\n'
    '                max_tokens=4096\n'
    '            )\n\n'
    '            # Check if agent is done\n'
    '            if response.stop_reason == "end_turn":\n'
    '                return self.extract_result(response)\n\n'
    '            # ACT: Execute any tool calls\n'
    '            tool_results = []\n'
    '            for block in response.content:\n'
    '                if block.type == "tool_use":\n'
    '                    result = await self.execute_tool(block)\n'
    '                    tool_results.append(result)\n\n'
    '            # OBSERVE: Feed results back for next reasoning step\n'
    '            messages.append({"role": "assistant", "content": response.content})\n'
    '            messages.append({"role": "user", "content": tool_results})'
)

add_heading_styled("5.3 Prospecting Agent — Deep Dive", 2)

add_body(
    "The Prospecting Agent is the most autonomous. Given the ICP criteria, it independently "
    "searches multiple sources, deduplicates, enriches, scores, and stores qualified leads."
)

add_code_block(
    'PROSPECTING AGENT SYSTEM PROMPT:\n\n'
    'You are AnalyticsGear\'s Prospecting Agent. Your job is to find\n'
    '25 qualified leads per day that match our Ideal Customer Profile.\n\n'
    'ICP Criteria:\n'
    '- Companies: 100-5000 employees, in Banking/Retail/Healthcare/SaaS/Manufacturing\n'
    '- Roles: CTO, VP Engineering, Head of Data, CDO, Data Engineering Manager\n'
    '- Signals: Uses Snowflake/Databricks/BigQuery, hiring data roles, recent funding\n\n'
    'Your workflow:\n'
    '1. Search Apollo.io for companies matching ICP (use apollo_search tool)\n'
    '2. Search LinkedIn for decision makers at those companies (use linkedin_search tool)\n'
    '3. Check job boards for companies hiring data roles (use job_scraper tool)\n'
    '4. For each lead, check if they already exist in CRM (use crm_lookup tool)\n'
    '5. Score new leads 1-100 based on ICP fit (use your judgment)\n'
    '6. Store qualified leads (score >= 25) in the vector DB (use vector_store tool)\n'
    '7. Log results to CRM (use crm_write tool)\n\n'
    'Be thorough but efficient. Skip leads that are clearly outside ICP.\n'
    'Report: how many found, how many qualified, top 5 hottest leads.'
)

add_heading_styled("5.4 Multi-Agent Coordination", 2)

add_body(
    "The Orchestrator manages agent execution order and passes context between agents:"
)

add_code_block(
    'class PipelineOrchestrator:\n'
    '    """Coordinates the daily pipeline across all agents."""\n\n'
    '    async def run_daily_pipeline(self):\n'
    '        # Phase 1: Prospecting (6-7 AM)\n'
    '        new_leads = await self.prospecting_agent.run(\n'
    '            "Find 25 new qualified leads matching our ICP"\n'
    '        )\n'
    '        log(f"Prospecting complete: {new_leads[\'count\']} leads found")\n\n'
    '        # Phase 2: Outreach (8 AM)\n'
    '        outreach_results = await self.outreach_agent.run(\n'
    '            f"Send personalized outreach to these leads: {new_leads[\'hot_leads\']}"\n'
    '            f"Also continue sequences for existing leads in pipeline"\n'
    '        )\n'
    '        log(f"Outreach complete: {outreach_results[\'emails_sent\']} sent")\n\n'
    '        # Phase 3: Follow-up (9 AM)\n'
    '        followup_results = await self.followup_agent.run(\n'
    '            "Check for new replies, engagement signals, and LinkedIn "\n'
    '            "acceptances. Draft responses and update lead statuses."\n'
    '        )\n'
    '        log(f"Follow-up complete: {followup_results[\'replies_handled\']}")\n\n'
    '        # Phase 4: Daily summary\n'
    '        await self.analytics_agent.run(\n'
    '            "Generate today\'s pipeline summary and send to founders"\n'
    '        )'
)

add_page_break()

# ════════════════════════════════════════════════════════════════════
#  6. COMPONENT 3 — MCP INTEGRATION LAYER
# ════════════════════════════════════════════════════════════════════

add_heading_styled("6. Component 3 — MCP Integration Layer", 1)

add_learning_badge("MCP", "Building MCP servers, tool schemas, resource endpoints, transport protocols, server composition")

add_body(
    "MCP servers are the connective tissue between your AI agents and the outside world. "
    "Each server exposes a set of tools that agents can discover and call. This is where "
    "you learn MCP by building real, production servers."
)

add_heading_styled("6.1 MCP Servers to Build", 2)

add_styled_table(
    ["MCP Server", "Tools Exposed", "What You Learn"],
    [
        ["mcp-apollo", "search_companies, search_contacts, get_company_details, verify_email",
         "REST API wrapping, pagination handling, rate limiting"],
        ["mcp-linkedin", "search_people, send_connection, send_message, get_profile, get_posts",
         "OAuth flows, scraping fallbacks, anti-detection"],
        ["mcp-email", "send_email, check_deliverability, get_opens_clicks, get_replies",
         "SMTP/API integration, tracking pixels, webhook handling"],
        ["mcp-crm", "read_leads, write_lead, update_status, log_activity, get_pipeline_stats",
         "Google Sheets API, schema management, concurrent writes"],
        ["mcp-vectordb", "store_embedding, search_similar, get_by_id, update_metadata, delete",
         "Embedding pipelines, vector operations, index management"],
        ["mcp-scraper", "scrape_job_boards, scrape_company_page, get_tech_stack, get_news",
         "Web scraping, HTML parsing, anti-bot handling"],
    ],
    col_widths=[3, 6, 7.5],
)

add_heading_styled("6.2 MCP Server Implementation Example", 2)

add_body("Here's how to build the CRM MCP server that connects agents to Google Sheets:")

add_code_block(
    'from mcp.server import Server\n'
    'from mcp.types import Tool, TextContent\n'
    'import gspread\n\n'
    'server = Server("mcp-crm")\n\n'
    '# Authenticate with Google Sheets\n'
    'gc = gspread.service_account(filename="credentials.json")\n'
    'sheet = gc.open("AnalyticsGear Sales Pipeline")\n\n'
    '@server.list_tools()\n'
    'async def list_tools():\n'
    '    return [\n'
    '        Tool(\n'
    '            name="read_leads",\n'
    '            description="Read leads from CRM filtered by status or tier",\n'
    '            inputSchema={\n'
    '                "type": "object",\n'
    '                "properties": {\n'
    '                    "status": {"type": "string", "enum": ["new","contacted","replied"]},\n'
    '                    "tier": {"type": "string", "enum": ["HOT","WARM","COLD"]},\n'
    '                    "limit": {"type": "integer", "default": 50}\n'
    '                }\n'
    '            }\n'
    '        ),\n'
    '        Tool(\n'
    '            name="write_lead",\n'
    '            description="Add a new lead to the CRM",\n'
    '            inputSchema={\n'
    '                "type": "object",\n'
    '                "properties": {\n'
    '                    "name": {"type": "string"},\n'
    '                    "email": {"type": "string"},\n'
    '                    "company": {"type": "string"},\n'
    '                    "title": {"type": "string"},\n'
    '                    "score": {"type": "integer"},\n'
    '                    "tier": {"type": "string"},\n'
    '                    "source": {"type": "string"}\n'
    '                },\n'
    '                "required": ["name", "email", "company"]\n'
    '            }\n'
    '        ),\n'
    '        Tool(\n'
    '            name="log_activity",\n'
    '            description="Log a sales activity (email sent, call made, etc.)",\n'
    '            inputSchema={\n'
    '                "type": "object",\n'
    '                "properties": {\n'
    '                    "lead_id": {"type": "string"},\n'
    '                    "channel": {"type": "string"},\n'
    '                    "action": {"type": "string"},\n'
    '                    "details": {"type": "string"}\n'
    '                },\n'
    '                "required": ["lead_id", "channel", "action"]\n'
    '            }\n'
    '        )\n'
    '    ]\n\n'
    '@server.call_tool()\n'
    'async def call_tool(name: str, arguments: dict):\n'
    '    if name == "write_lead":\n'
    '        ws = sheet.worksheet("Master Leads")\n'
    '        ws.append_row([\n'
    '            generate_id(), arguments["name"], arguments["email"],\n'
    '            arguments.get("company"), arguments.get("title"),\n'
    '            arguments.get("score", 0), arguments.get("tier", "COLD"),\n'
    '            arguments.get("source", "unknown"), datetime.now().isoformat()\n'
    '        ])\n'
    '        return [TextContent(type="text", text=f"Lead {arguments[\'name\']} added")]'
)

add_heading_styled("6.3 Why MCP Matters (For You and Your Clients)", 2)

add_callout_box(
    "Without MCP: Every AI integration is custom code. Adding HubSpot means rewriting agent logic.\n"
    "With MCP: Adding HubSpot is just spinning up a new MCP server. Zero agent changes.\n\n"
    "This is exactly what enterprise clients need — and you'll be able to build it because you've "
    "done it for your own pipeline. MCP server development is a high-value, low-competition skill "
    "right now (April 2026). Most consultancies are still writing custom API wrappers.",
    title="The Business Value of MCP:"
)

add_page_break()

# ════════════════════════════════════════════════════════════════════
#  7. COMPONENT 4 — CONVERSATION MEMORY
# ════════════════════════════════════════════════════════════════════

add_heading_styled("7. Component 4 — Intelligent Conversation Memory", 1)

add_learning_badge("RAG + Agents", "Long-term memory, conversation summarization, context window management, memory retrieval")

add_body(
    "Unlike simple CRM notes, the memory system gives agents persistent, searchable memory "
    "of every interaction. When the Outreach Agent contacts a prospect for the third time, "
    "it remembers everything — what was said, what they clicked, what objection they raised, "
    "and what worked with similar prospects."
)

add_heading_styled("7.1 Memory Layers", 2)

add_styled_table(
    ["Memory Layer", "What It Stores", "How It's Used"],
    [
        ["Lead Memory", "Every email sent/received, LinkedIn messages, call notes, engagement signals per lead",
         "Agents retrieve full interaction history before any touchpoint — no prospect hears the same pitch twice"],
        ["Pattern Memory", "Which subject lines get opens, which CTAs get clicks, which industries respond best",
         "The Outreach Agent learns from aggregate success patterns and improves its messaging"],
        ["Relationship Memory", "Tone preferences, objections raised, internal champions identified, referral connections",
         "Enables genuinely personal follow-ups — 'Last time you mentioned Q3 budget cycles...'"],
        ["Company Memory", "Evolving company profile — news, funding, hiring changes, tech stack updates",
         "Keeps outreach relevant as companies change — catches signals like new CTO or fresh funding"],
    ],
    col_widths=[3.5, 6, 7],
)

add_heading_styled("7.2 Memory-Augmented Outreach (Example)", 2)

add_code_block(
    'AGENT INTERNAL REASONING (FOLLOW-UP EMAIL #3 TO SARAH CHEN):\n\n'
    '  [Retrieve Memory] --> Query: "All interactions with Sarah Chen, DataFlow Inc"\n\n'
    '  [Memory Returns]:\n'
    '  - Email #1 (Mar 5): Sent data pipeline optimization pitch. Opened 3x, no reply.\n'
    '  - Email #2 (Mar 10): Shared Lakehouse migration blog post. Clicked link, read 4 min.\n'
    '  - LinkedIn: Connected Mar 7. She posted about "Airflow pain" on Mar 12.\n'
    '  - Company: DataFlow raised Series B ($18M) on Mar 1. Hiring 3 data engineers.\n\n'
    '  [Agent Thinks]:\n'
    '  "She engaged with the Lakehouse content and is posting about Airflow issues.\n'
    '   She didn\'t reply to a generic pitch but clicked a technical article.\n'
    '   They just raised funding and are hiring — they have budget and need.\n'
    '   I should reference Airflow specifically and offer something tactical, not salesy."\n\n'
    '  [Agent Writes]:\n'
    '  Subject: airflow at dataflow\n\n'
    '  Hi Sarah,\n\n'
    '  Saw your post about Airflow scheduling issues — we hit the exact same wall\n'
    '  with a client last quarter. Ended up moving their critical DAGs to Dagster\n'
    '  while keeping Airflow for batch jobs. Cut incident rate by 70%.\n\n'
    '  Happy to walk through the migration pattern if useful — 15 min, no pitch.\n\n'
    '  (This is the kind of email that gets replies.)'
)

add_page_break()

# ════════════════════════════════════════════════════════════════════
#  8. COMPONENT 5 — SELF-IMPROVING FEEDBACK LOOP
# ════════════════════════════════════════════════════════════════════

add_heading_styled("8. Component 5 — Self-Improving Feedback Loop", 1)

add_learning_badge("Evaluation + RLHF", "LLM evaluation, A/B testing, reinforcement from human feedback, metric-driven optimization")

add_body(
    "This is what separates a toy project from a production AI system. The pipeline doesn't just "
    "run — it learns. Every email opened, every reply received, every meeting booked becomes "
    "training data that makes the system better."
)

add_heading_styled("8.1 The Feedback Loop", 2)

add_code_block(
    "CONTINUOUS IMPROVEMENT CYCLE:\n\n"
    "  [Send Outreach]  -->  [Track Outcomes]  -->  [Analyze Patterns]\n"
    "        ^                                            |\n"
    "        |                                            v\n"
    "  [Update Prompts  <--  [Store Learnings]  <--  [AI Identifies\n"
    "   & Strategies]        [in Vector DB]          What Works]"
)

add_heading_styled("8.2 What Gets Measured and Fed Back", 2)

add_styled_table(
    ["Signal", "What The AI Learns", "How It Adapts"],
    [
        ["Email opened but no reply", "Subject line worked, but body or CTA didn't", "Test different CTAs, shorter body, different value props"],
        ["Email reply (positive)", "This messaging + context combination works", "Store as successful pattern, weight similar approaches higher"],
        ["Email reply (objection)", "Prospect has budget/timing/vendor concern", "Add objection to lead memory, adjust follow-up strategy"],
        ["LinkedIn post engagement", "Prospect is active and interested in topic X", "Reference their post in next outreach, align messaging to topic X"],
        ["Meeting booked", "Full sequence worked — from first touch to booking", "Analyze and replicate: which source, message style, timing, channel mix"],
        ["Meeting no-show", "Prospect was interested but not committed enough", "Add pre-meeting confirmation step, send value-add before meeting"],
        ["Deal won", "Everything worked — mark as gold-standard pattern", "Use this as a template for similar industry/role/company size combos"],
        ["Deal lost", "Something failed in the process", "Post-mortem analysis: was it timing, pricing, wrong champion, or poor fit?"],
    ],
    col_widths=[3, 5.5, 8],
)

add_heading_styled("8.3 Weekly Self-Improvement Cycle", 2)

add_numbered([
    "Analytics Agent pulls all engagement data from the past 7 days",
    "It identifies top-performing and worst-performing outreach by open rate, reply rate, and meeting conversion",
    "It compares messaging patterns: which industries respond to which angles, which subject line styles work",
    "It generates specific recommendations: 'Emails mentioning cost reduction have 3x reply rate in banking vertical'",
    "Successful email/message patterns are embedded and stored in the 'successful_patterns' vector collection",
    "The Outreach Agent's RAG retrieval now surfaces these winning patterns when crafting new outreach",
    "Lead scoring weights are adjusted based on which lead profiles actually convert to meetings",
    "A weekly report summarizes all learnings and is sent to the founders for review",
])

add_page_break()

# ════════════════════════════════════════════════════════════════════
#  9. TECH STACK
# ════════════════════════════════════════════════════════════════════

add_heading_styled("9. Tech Stack & Infrastructure", 1)

add_heading_styled("9.1 Core Stack", 2)

add_styled_table(
    ["Layer", "Technology", "Why This Choice"],
    [
        ["Language", "Python 3.12+", "Best AI/ML ecosystem, all SDKs available"],
        ["LLM Provider", "Anthropic Claude API (Sonnet for agents, Haiku for scoring)", "Best tool use, MCP creator, cost-effective"],
        ["Vector Database", "Qdrant (local dev) → Qdrant Cloud (prod)", "Rich filtering, fast, Python-native, free self-host"],
        ["Embeddings", "Voyage AI (voyage-3) or OpenAI text-embedding-3-small", "High quality, good price-performance"],
        ["Agent Framework", "Claude Agent SDK or custom ReAct loop", "Native MCP support, clean abstraction"],
        ["MCP Framework", "MCP Python SDK (@modelcontextprotocol/sdk)", "Official SDK, well-documented"],
        ["Orchestration", "Prefect or APScheduler", "Scheduling, retries, observability"],
        ["Email", "SendGrid API + Instantly.ai for warmup", "Reliable delivery, tracking, warmup"],
        ["CRM", "Google Sheets API (gspread) → HubSpot API later", "Free start, easy migration path"],
        ["Monitoring", "Langfuse or LangSmith", "LLM observability — trace costs, latency, quality"],
        ["Deployment", "Docker on a small VPS (Hetzner/DigitalOcean)", "Cheap, full control, runs daily"],
    ],
    col_widths=[3, 5.5, 8],
)

add_heading_styled("9.2 Project Structure", 2)

add_code_block(
    'sales-pipeline-ai/\n'
    '├── agents/\n'
    '│   ├── base_agent.py              # ReAct agent base class\n'
    '│   ├── prospecting_agent.py       # Lead discovery agent\n'
    '│   ├── outreach_agent.py          # Email & LinkedIn outreach agent\n'
    '│   ├── followup_agent.py          # Engagement monitoring agent\n'
    '│   ├── analytics_agent.py         # Performance analysis agent\n'
    '│   └── orchestrator.py            # Multi-agent coordinator\n'
    '├── mcp_servers/\n'
    '│   ├── mcp_apollo/                # Apollo.io MCP server\n'
    '│   │   ├── server.py\n'
    '│   │   └── tools.py\n'
    '│   ├── mcp_linkedin/              # LinkedIn MCP server\n'
    '│   ├── mcp_email/                 # SendGrid email MCP server\n'
    '│   ├── mcp_crm/                   # Google Sheets CRM MCP server\n'
    '│   ├── mcp_vectordb/              # Qdrant vector DB MCP server\n'
    '│   └── mcp_scraper/               # Web scraping MCP server\n'
    '├── rag/\n'
    '│   ├── embedder.py                # Embedding pipeline\n'
    '│   ├── chunker.py                 # Document chunking strategies\n'
    '│   ├── retriever.py               # RAG retrieval logic\n'
    '│   └── collections.py             # Vector DB collection schemas\n'
    '├── memory/\n'
    '│   ├── lead_memory.py             # Per-lead interaction memory\n'
    '│   ├── pattern_memory.py          # Successful outreach patterns\n'
    '│   └── memory_manager.py          # Memory storage & retrieval\n'
    '├── feedback/\n'
    '│   ├── tracker.py                 # Engagement signal tracking\n'
    '│   ├── analyzer.py                # Pattern analysis\n'
    '│   └── optimizer.py               # Strategy optimization\n'
    '├── config/\n'
    '│   ├── icp.yaml                   # Ideal customer profile\n'
    '│   ├── sequences.yaml             # Outreach sequences\n'
    '│   └── scoring.yaml               # Lead scoring criteria\n'
    '├── prompts/\n'
    '│   ├── prospecting.txt            # Prospecting agent system prompt\n'
    '│   ├── outreach.txt               # Outreach agent system prompt\n'
    '│   ├── followup.txt               # Follow-up agent system prompt\n'
    '│   └── analytics.txt              # Analytics agent system prompt\n'
    '├── data/\n'
    '│   ├── qdrant_data/               # Local vector DB storage\n'
    '│   └── target_companies.csv       # Initial target list\n'
    '├── tests/\n'
    '├── docker-compose.yml             # Qdrant + app containers\n'
    '├── Dockerfile\n'
    '├── .env                           # API keys\n'
    '├── requirements.txt\n'
    '└── README.md'
)

add_page_break()

# ════════════════════════════════════════════════════════════════════
#  10. BUILD & LEARN ROADMAP
# ════════════════════════════════════════════════════════════════════

add_heading_styled("10. Build-and-Learn Roadmap (12 Weeks)", 1)

add_body(
    "This roadmap is structured so each phase teaches specific AI concepts while building "
    "production components. By week 12, you have both a working pipeline and deep expertise."
)

# Phase 1
add_heading_styled("Phase 1: Foundation (Weeks 1-3)", 2)

add_callout_box(
    "Vector Databases, Embeddings, Basic RAG, Claude API tool use",
    title="AI Skills You'll Learn:"
)

add_styled_table(
    ["Week", "Build", "Learn"],
    [
        ["Week 1", "Set up Qdrant locally. Build the embedding pipeline. Ingest 100 company profiles and all AnalyticsGear blog content into vector collections.",
         "How embeddings work, chunking strategies, similarity search, metadata filtering. Play with different embedding models and compare quality."],
        ["Week 2", "Build the RAG retrieval system. Given a prospect, retrieve relevant company data + matching AG content. Test retrieval quality.",
         "RAG architecture, prompt injection with retrieved context, retrieval evaluation (are the right docs coming back?), tuning top-K and similarity thresholds."],
        ["Week 3", "Build lead scoring with Claude API tool use. The LLM calls a 'score_lead' tool that evaluates prospects against ICP criteria.",
         "Claude API tool use, structured output, JSON mode, prompt engineering for consistent scoring, handling edge cases."],
    ],
    col_widths=[1.5, 7.5, 7.5],
)

# Phase 2
add_heading_styled("Phase 2: MCP Servers (Weeks 4-6)", 2)

add_callout_box(
    "MCP protocol, server development, tool schemas, transport layers, server composition",
    title="AI Skills You'll Learn:"
)

add_styled_table(
    ["Week", "Build", "Learn"],
    [
        ["Week 4", "Build mcp-crm server (Google Sheets). Build mcp-vectordb server (Qdrant). Test both servers with Claude Desktop.",
         "MCP protocol basics, tool schema design, server lifecycle, stdio transport, testing MCP servers locally."],
        ["Week 5", "Build mcp-email server (SendGrid). Build mcp-scraper server (job boards, company pages). Set up email warmup.",
         "REST API wrapping into MCP, handling async operations, error handling in MCP servers, webhook integration."],
        ["Week 6", "Build mcp-apollo server and mcp-linkedin server. Connect all servers and test tool discovery.",
         "OAuth in MCP, rate limiting, anti-detection strategies, composing multiple MCP servers for an agent."],
    ],
    col_widths=[1.5, 7.5, 7.5],
)

# Phase 3
add_heading_styled("Phase 3: Agents (Weeks 7-9)", 2)

add_callout_box(
    "Agent architecture, ReAct loop, multi-agent systems, planning, error recovery, tool orchestration",
    title="AI Skills You'll Learn:"
)

add_styled_table(
    ["Week", "Build", "Learn"],
    [
        ["Week 7", "Build the base agent class with ReAct loop. Build the Prospecting Agent. Test it finding and scoring 25 leads end-to-end.",
         "ReAct pattern, agent loop design, tool selection by LLM, handling tool errors, agent observability (logging each step)."],
        ["Week 8", "Build the Outreach Agent with RAG-powered personalization. It retrieves context and generates + sends emails.",
         "RAG-augmented agents, multi-tool orchestration, the agent deciding WHICH tools to use and in what order."],
        ["Week 9", "Build Follow-Up Agent and Analytics Agent. Build the Orchestrator that coordinates all four agents on a daily schedule.",
         "Multi-agent coordination, shared state management, agent handoffs, scheduling, error recovery across agents."],
    ],
    col_widths=[1.5, 7.5, 7.5],
)

# Phase 4
add_heading_styled("Phase 4: Memory & Feedback (Weeks 10-12)", 2)

add_callout_box(
    "Long-term memory, conversation history, feedback loops, LLM evaluation, self-improvement, production deployment",
    title="AI Skills You'll Learn:"
)

add_styled_table(
    ["Week", "Build", "Learn"],
    [
        ["Week 10", "Build the memory system — per-lead interaction history, pattern memory for successful outreach, relationship memory.",
         "Memory architecture for agents, summarization for long histories, memory retrieval strategies, avoiding context window overflow."],
        ["Week 11", "Build the feedback loop — track engagement signals, analyze patterns, update successful_patterns collection, adjust scoring.",
         "LLM-as-judge evaluation, A/B testing frameworks, reinforcement from outcomes, metric-driven prompt optimization."],
        ["Week 12", "Deploy to production (Docker on VPS). Set up monitoring with Langfuse. Run full pipeline for a week. Fix issues.",
         "Production deployment, LLM observability, cost tracking, latency optimization, error alerting, the reality of production AI."],
    ],
    col_widths=[1.5, 7.5, 7.5],
)

add_page_break()

# ════════════════════════════════════════════════════════════════════
#  11. TURNING THIS INTO A CASE STUDY
# ════════════════════════════════════════════════════════════════════

add_heading_styled("11. Turning This Into a Case Study", 1)

add_body(
    "Once the pipeline is running and producing results, you package it as a case study "
    "that demonstrates every AI capability your clients are asking for. This is your strongest "
    "sales asset — you built it, you run it, and you have the numbers to prove it works."
)

add_heading_styled("11.1 Case Study Structure", 2)

add_callout_box(
    'TITLE: "How We Built an AI-Native Sales Pipeline That Books 15+ Meetings/Month"\n\n'
    "SUBTITLE: Using Multi-Agent Architecture, RAG, Vector Databases, and MCP\n\n"
    "SECTIONS:\n"
    "1. The Challenge — Why we needed automated, intelligent sales outreach\n"
    "2. The Architecture — Multi-agent system with RAG-powered personalization\n"
    "3. The Tech — MCP servers, Qdrant, Claude API, embedding pipelines\n"
    "4. The Results — Meetings booked, reply rates, cost savings, system uptime\n"
    "5. Key Learnings — What worked, what didn't, how the system improved itself\n"
    "6. How This Applies to Your Business — Bridge to client's use case",
    title="Case Study Template:"
)

add_heading_styled("11.2 Metrics to Showcase", 2)

add_styled_table(
    ["Metric Category", "Specific Metrics to Highlight"],
    [
        ["Scale", "Leads processed per month, emails sent, companies analyzed"],
        ["AI Performance", "RAG retrieval accuracy, agent task completion rate, scoring accuracy vs. actual conversion"],
        ["Business Results", "Reply rate, meetings booked, pipeline value generated, cost per meeting"],
        ["System Reliability", "Uptime, error rate, daily processing time, recovery from failures"],
        ["Self-Improvement", "Week-over-week improvement in reply rates, how messaging evolved based on feedback"],
        ["Cost Efficiency", "Total AI spend vs. results, comparison to human SDR cost, ROI"],
    ],
    col_widths=[4, 12.5],
)

add_heading_styled("11.3 Before/After Comparison", 2)

add_styled_table(
    ["Aspect", "Before (Manual/Scripts)", "After (AI-Native Pipeline)"],
    [
        ["Personalization", "Generic templates with {name} merge tags", "RAG-powered emails referencing company's tech stack, recent news, and pain points"],
        ["Lead Scoring", "Manual gut feeling or simple rules", "AI scoring on 5 dimensions with continuous improvement from conversion data"],
        ["Memory", "Spreadsheet notes, no context across touchpoints", "Vector-stored interaction history, agents remember everything"],
        ["Tool Integration", "Custom API code per integration, brittle", "MCP servers — add new tools without changing agent code"],
        ["Improvement", "Manual A/B testing, quarterly reviews", "Automated feedback loop, weekly self-optimization"],
        ["Scale", "30 emails/day, one person's bandwidth", "200+ personalized touchpoints/day across email + LinkedIn"],
    ],
    col_widths=[3, 6, 7.5],
)

add_page_break()

# ════════════════════════════════════════════════════════════════════
#  12. PORTFOLIO PRESENTATION GUIDE
# ════════════════════════════════════════════════════════════════════

add_heading_styled("12. Portfolio Presentation Guide", 1)

add_body(
    "When presenting to potential clients, tailor the case study to their use case. "
    "The pipeline maps cleanly to problems every data-driven company faces."
)

add_heading_styled("12.1 Mapping Pipeline Components to Client Use Cases", 2)

add_styled_table(
    ["Pipeline Component", "Maps To (Client Use Case)", "Industries"],
    [
        ["RAG + Vector DB for company profiles", "Customer 360 views, product search, knowledge bases, support bots",
         "Retail, SaaS, Banking, Healthcare"],
        ["AI Agents with tool use", "Automated workflows, document processing, internal copilots",
         "All industries"],
        ["MCP Integration Layer", "Enterprise AI integration, connecting LLMs to internal systems",
         "Enterprise, Banking, Manufacturing"],
        ["Conversation Memory", "Customer service memory, patient history, case management",
         "Healthcare, Banking, SaaS"],
        ["Feedback Loop / Self-Improvement", "Recommendation systems, fraud detection, quality prediction",
         "Retail, Banking, Manufacturing"],
        ["Multi-Agent Orchestration", "Complex business process automation, supply chain optimization",
         "Logistics, Manufacturing, Finance"],
    ],
    col_widths=[4, 6.5, 6],
)

add_heading_styled("12.2 Pitch Angles by Client Type", 2)

pitches = [
    ("For CTOs / VP Engineering",
     '"We built a multi-agent system with 6 MCP tool integrations and a RAG pipeline processing '
     '500+ documents. It runs in production daily with 99.5% uptime. We can build the same '
     'architecture for your [use case]."'),
    ("For Heads of Data",
     '"Our vector database handles 10,000+ company embeddings with sub-200ms retrieval. '
     'The embedding pipeline processes new data daily and the RAG system maintains retrieval '
     'accuracy above 90%. We can apply this to your data products."'),
    ("For Business Leaders",
     '"We replaced a $60K/year sales hire with an AI system that costs $4K/year and books '
     '15+ qualified meetings per month. It improved its own performance by 40% over 3 months '
     'through automated learning. What process in your business could benefit from this?"'),
]

for title, pitch in pitches:
    add_heading_styled(title, 3)
    add_body(pitch)

add_heading_styled("12.3 Content to Create From This Project", 2)

add_numbered([
    "Detailed case study (PDF / blog post) with architecture diagrams and real metrics",
    "Technical blog series: 'Building a Multi-Agent Sales Pipeline' (4-5 posts for AnalyticsGear Insights)",
    "Open-source the MCP servers (mcp-crm, mcp-vectordb) on GitHub — builds credibility and visibility",
    "Conference talk: 'RAG in Production: Lessons from Building an AI Sales Pipeline' (submit to data/AI conferences)",
    "LinkedIn post series documenting the build journey — attracts exactly the audience you want to sell to",
    "Live demo in sales calls — show the actual pipeline running, explain each component",
    "YouTube walkthrough or webinar — 'How We Built an AI-Native Sales Pipeline from Scratch'",
])

add_page_break()

# ════════════════════════════════════════════════════════════════════
#  13. WHAT YOU'LL BE ABLE TO SELL
# ════════════════════════════════════════════════════════════════════

add_heading_styled("13. What You'll Be Able to Sell After Building This", 1)

add_body(
    "Building this pipeline doesn't just give you a case study — it gives you productized "
    "service offerings backed by real production experience."
)

add_styled_table(
    ["Service Offering", "What You Deliver", "Your Credibility", "Price Range"],
    [
        ["RAG System Development", "Production RAG pipelines for knowledge bases, support bots, search, and copilots",
         "You run one daily — 500+ docs, sub-200ms retrieval", "$15K - $50K"],
        ["AI Agent Development", "Custom autonomous agents for business process automation",
         "You have 4 agents in production with multi-tool orchestration", "$20K - $75K"],
        ["MCP Server Development", "Connect enterprise LLMs to internal tools and systems via MCP",
         "You've built 6 production MCP servers", "$10K - $30K per integration"],
        ["Vector Database Implementation", "Design and deploy vector search for product discovery, recommendations, similarity",
         "You manage a 10K+ vector collection in production", "$15K - $40K"],
        ["AI Pipeline Architecture", "End-to-end AI system design — from data ingestion to LLM output to feedback loops",
         "Your entire sales pipeline is a reference architecture", "$25K - $100K"],
        ["LLM Evaluation & Optimization", "Set up evaluation frameworks, A/B testing, cost optimization for LLM applications",
         "Your pipeline self-optimized and improved 40% over 3 months", "$10K - $25K"],
    ],
    col_widths=[3.5, 5, 4.5, 3.5],
)

add_callout_box(
    "Combined annual revenue potential from these new service lines: $200K - $500K+\n\n"
    "Total investment to build the pipeline: ~$4K in tools + 12 weeks of founder time\n\n"
    "The pipeline pays for itself with the first meeting it books. Every service you sell "
    "afterwards is pure upside enabled by the expertise you built.",
    title="ROI of Building This:"
)

add_page_break()

# ════════════════════════════════════════════════════════════════════
#  QUICK START
# ════════════════════════════════════════════════════════════════════

add_heading_styled("Getting Started — This Week", 1)

steps = [
    ("Day 1", "Install Qdrant locally (Docker), get Claude API key, set up the project repo with the structure from Section 9.2"),
    ("Day 2", "Write your first embedding pipeline — embed 10 company profiles into Qdrant and run similarity searches"),
    ("Day 3", "Build a basic RAG flow — given a prospect name, retrieve their profile + a relevant AG blog post, generate a personalized email"),
    ("Day 4", "Add Claude tool use — let the LLM call a 'score_lead' tool and a 'retrieve_context' tool as part of generating outreach"),
    ("Day 5", "Build your first MCP server (mcp-crm for Google Sheets). Test it with Claude Desktop — see the magic of tool discovery"),
    ("Weekend", "Reflect on what you learned. Read the MCP spec. Plan Week 2: more MCP servers + the first agent"),
]

for i, (when, what) in enumerate(steps, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"  {i}. ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = PURPLE
    r2 = p.add_run(f"[{when}]  ")
    r2.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = TEAL
    r3 = p.add_run(what)
    r3.font.size = Pt(10)
    r3.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(8)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("━" * 40)
r.font.color.rgb = PURPLE
r.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Build it. Learn it. Sell it.\n"
    "The pipeline is the product, the education, and the proof — all in one."
)
r.font.size = Pt(10)
r.font.color.rgb = GRAY
r.italic = True

# ── Save ──
output_path = r"c:\analyticsgear\sales_pipeline\AnalyticsGear_AI_Native_Pipeline_Technical_Blueprint.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
