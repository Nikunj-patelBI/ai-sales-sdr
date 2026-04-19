"""Generate a professional Word document for the AI Sales Pipeline Blueprint."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Color palette ──
INDIGO = RGBColor(0x4F, 0x46, 0xE5)       # Primary brand
DARK_INDIGO = RGBColor(0x31, 0x2E, 0x81)  # Headings
TEAL = RGBColor(0x0D, 0x94, 0x88)         # Accents
DARK = RGBColor(0x1E, 0x1E, 0x2E)         # Body text
GRAY = RGBColor(0x64, 0x64, 0x80)         # Subtle text
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF1, 0xF5, 0xF9)     # Table header bg hex: F1F5F9
TABLE_BORDER = RGBColor(0xE2, 0xE8, 0xF0)

# ── Style helpers ──
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def add_styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_INDIGO
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, "F1F5F9")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK
            if r_idx % 2 == 1:
                set_cell_shading(cell, "FAFAFA")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table

def add_heading_styled(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(22)
        run.font.color.rgb = DARK_INDIGO
        run.bold = True
        p.space_before = Pt(24)
        p.space_after = Pt(8)
        # Add a colored line below
        border_p = doc.add_paragraph()
        border_p.space_before = Pt(0)
        border_p.space_after = Pt(12)
        pPr = border_p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="4F46E5"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = INDIGO
        run.bold = True
        p.space_before = Pt(18)
        p.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = TEAL
        run.bold = True
        p.space_before = Pt(12)
        p.space_after = Pt(4)
    return p

def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    p.paragraph_format.line_spacing = Pt(16)
    p.space_after = Pt(6)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = DARK
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK
    else:
        p.text = ""
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(2)
    return p

def add_callout_box(text, title=None):
    """Add a highlighted callout box using a single-cell table."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "EEF2FF")  # Light indigo background

    if title:
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = DARK_INDIGO
        p2 = cell.add_paragraph()
        r2 = p2.add_run(text)
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK
    else:
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = DARK

    # Border styling
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="4F46E5"/>'
        '  <w:left w:val="single" w:sz="12" w:space="0" w:color="4F46E5"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="4F46E5"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="4F46E5"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    doc.add_paragraph()
    return table

def add_numbered(items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = INDIGO
        r2 = p.add_run(item)
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(2)

def add_page_break():
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════

# Spacer
for _ in range(4):
    doc.add_paragraph()

# Company name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ANALYTICSGEAR")
r.font.size = Pt(14)
r.font.color.rgb = INDIGO
r.bold = True
r.font.letter_spacing = Pt(4)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AI-Powered Sales Pipeline")
r.font.size = Pt(36)
r.font.color.rgb = DARK_INDIGO
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Blueprint & Implementation Guide")
r.font.size = Pt(20)
r.font.color.rgb = GRAY

# Decorative line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("━" * 40)
r.font.color.rgb = INDIGO
r.font.size = Pt(12)

doc.add_paragraph()

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("End-to-end automated prospecting, outreach, and pipeline management\ndesigned for AnalyticsGear's Data, AI & Cloud consulting practice")
r.font.size = Pt(11)
r.font.color.rgb = GRAY
r.italic = True

for _ in range(4):
    doc.add_paragraph()

# Meta info
meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Document Version", "1.0"),
    ("Date", datetime.date.today().strftime("%B %d, %Y")),
    ("Prepared For", "AnalyticsGear Leadership"),
    ("Classification", "Internal / Confidential"),
]
for i, (label, value) in enumerate(meta_data):
    c1 = meta_table.rows[i].cells[0]
    c2 = meta_table.rows[i].cells[1]
    c1.text = ""
    c2.text = ""
    r1 = c1.paragraphs[0].add_run(label)
    r1.font.size = Pt(9)
    r1.font.color.rgb = GRAY
    r1.bold = True
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = c2.paragraphs[0].add_run(value)
    r2.font.size = Pt(9)
    r2.font.color.rgb = DARK
    c1.width = Cm(5)
    c2.width = Cm(6)

add_page_break()

# ════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════

add_heading_styled("Table of Contents", 1)

toc_items = [
    "Executive Summary",
    "Pipeline Architecture Overview",
    "Phase 1 — Lead Discovery & Prospecting",
    "Phase 2 — Lead Enrichment & Scoring",
    "Phase 3 — Outreach Automation",
    "Phase 4 — Follow-Up & Nurturing",
    "Phase 5 — CRM, Logging & Tracking",
    "Phase 6 — Analytics & Optimization",
    "Tech Stack & Tools",
    "Daily Automation Schedule",
    "Implementation Roadmap",
    "Cost Estimates",
    "Risk Mitigation & Compliance",
    "Email & LinkedIn Templates",
]

for i, item in enumerate(toc_items, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"  {i:02d}   ")
    r.font.size = Pt(10)
    r.font.color.rgb = INDIGO
    r.bold = True
    r2 = p.add_run(item)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)

add_page_break()

# ════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════

add_heading_styled("1. Executive Summary", 1)

add_body(
    "AnalyticsGear offers Data Engineering, Analytics, AI, Cloud, and DevOps consulting "
    "services to mid-to-large enterprises across Banking, Retail, Healthcare, Manufacturing, "
    "Logistics, and SaaS verticals."
)

add_body(
    "Rather than hiring a sales team upfront, this blueprint designs a fully automated AI "
    "sales pipeline that runs 2-4 hours per day with minimal human oversight, handling the "
    "work of a full-time sales development representative at approximately 10% of the cost."
)

add_callout_box(
    "Discover 50-100 new qualified leads per week from multiple sources\n"
    "Enrich each lead with company data, tech stack, hiring signals, and pain points\n"
    "Reach out via personalized cold email, LinkedIn, and scheduled cold calls\n"
    "Follow up intelligently based on engagement signals\n"
    "Log everything in a central CRM with full audit trail\n"
    "Run 2-4 hours/day on a scheduled basis with minimal human oversight",
    title="What the Pipeline Does:"
)

add_body(
    "Expected outcome: 10-20 qualified conversations per month within 60 days of launch, "
    "with a fully traceable pipeline from first touch to booked discovery call."
)

add_page_break()

# ════════════════════════════════════════════════════════════════
#  2. PIPELINE ARCHITECTURE OVERVIEW
# ════════════════════════════════════════════════════════════════

add_heading_styled("2. Pipeline Architecture Overview", 1)

add_body(
    "The pipeline operates in six sequential phases, running daily on an automated schedule. "
    "Each phase feeds data into the next, with a central CRM layer capturing all activities "
    "and an analytics layer providing continuous optimization insights."
)

# Architecture flow table
arch_table = doc.add_table(rows=2, cols=4)
arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(arch_table)

phases_top = [
    ("PHASE 1", "Lead Discovery\n& Prospecting"),
    ("PHASE 2", "Lead Enrichment\n& Scoring"),
    ("PHASE 3", "Outreach\nAutomation"),
    ("PHASE 4", "Follow-Up\n& Nurturing"),
]

for i, (phase, desc) in enumerate(phases_top):
    cell = arch_table.rows[0].cells[i]
    set_cell_shading(cell, "EEF2FF")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(phase)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = INDIGO

    cell2 = arch_table.rows[1].cells[i]
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(desc)
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = DARK

doc.add_paragraph()

# Bottom phases
arch_table2 = doc.add_table(rows=2, cols=2)
arch_table2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(arch_table2)

for i, (phase, desc) in enumerate([
    ("PHASE 5", "CRM, Logging & Tracking"),
    ("PHASE 6", "Analytics & Optimization"),
]):
    cell = arch_table2.rows[0].cells[i]
    set_cell_shading(cell, "F0FDFA")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(phase)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = TEAL
    cell2 = arch_table2.rows[1].cells[i]
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(desc)
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = DARK

doc.add_paragraph()

add_body(
    "All phases log activities to a central CRM (Google Sheets, later HubSpot). "
    "The analytics layer runs weekly to generate performance reports and optimization recommendations."
)

add_page_break()

# ════════════════════════════════════════════════════════════════
#  3. PHASE 1 — LEAD DISCOVERY
# ════════════════════════════════════════════════════════════════

add_heading_styled("3. Phase 1 — Lead Discovery & Prospecting", 1)

add_heading_styled("3.1 Ideal Customer Profile (ICP)", 2)

add_body(
    "The foundation of effective prospecting is a well-defined Ideal Customer Profile. "
    "Based on AnalyticsGear's service offerings and past engagements, our ICP is defined as follows:"
)

add_styled_table(
    ["Attribute", "Criteria"],
    [
        ["Company Size", "100 - 5,000 employees"],
        ["Revenue", "$10M - $500M"],
        ["Industries", "Banking/Finance, Retail/CPG, Healthcare, Manufacturing, Logistics, SaaS"],
        ["Target Roles", "CTO, VP Engineering, Head of Data, Data Engineering Manager, VP Analytics, CDO"],
        ["Geography", "US, UK, EU, Middle East, India, Southeast Asia"],
        ["Tech Signals", "Uses Snowflake, Databricks, BigQuery, dbt, Airflow; hiring data engineers; recent funding"],
        ["Pain Signals", "Job postings for data roles, legacy migration mentions, compliance requirements"],
    ],
    col_widths=[4, 13],
)

add_heading_styled("3.2 Lead Sources & Collection Methods", 2)

sources = [
    ("Source 1: LinkedIn Sales Navigator", "Primary",
     "Search for ICP-matching decision makers using advanced filters. Export via Sales Navigator API or tools like Phantombuster/Apify. Target: 20-30 leads/day."),
    ("Source 2: Company Databases (Apollo.io)", "Primary",
     "Bulk search companies matching ICP criteria via API. Filter by industry, size, tech stack, and hiring signals. Target: 30-50 leads/day."),
    ("Source 3: Job Board Scraping", "Intent Signal",
     "Companies hiring data engineers, analytics managers, and ML engineers have active data needs and budget. Scrape LinkedIn Jobs, Indeed, and Glassdoor. HIGH signal strength. Target: 10-20 companies/day."),
    ("Source 4: Technology Signal Monitoring", "Intent Signal",
     "Track companies adopting or struggling with Snowflake, Databricks, dbt, Airflow via BuiltWith, Wappalyzer, and G2 reviews. Recent adoption = needs implementation help."),
    ("Source 5: Event & Conference Attendees", "Supplementary",
     "Scrape attendee lists from data/cloud/AI conferences: dbt Coalesce, Snowflake Summit, AWS re:Invent, DataEngBytes, local meetups."),
    ("Source 6: Inbound Website Leads", "Highest Priority",
     "Auto-forward Web3Forms contact form submissions to the pipeline for immediate follow-up. These leads are already warm and should be contacted within 1 hour."),
    ("Source 7: Social Listening", "Supplementary",
     'Monitor LinkedIn, Twitter/X, Reddit, and Hacker News for keywords like "data pipeline broken", "need data engineer", "migrating to snowflake". Engage and add to pipeline.'),
]

for title, priority, desc in sources:
    add_heading_styled(title, 3)
    p = doc.add_paragraph()
    r = p.add_run(f"Priority: {priority}")
    r.font.size = Pt(9)
    r.font.color.rgb = TEAL
    r.italic = True
    r.bold = True
    add_body(desc)

add_page_break()

# ════════════════════════════════════════════════════════════════
#  4. PHASE 2 — ENRICHMENT & SCORING
# ════════════════════════════════════════════════════════════════

add_heading_styled("4. Phase 2 — Lead Enrichment & Scoring", 1)

add_heading_styled("4.1 Enrichment Data Points", 2)

add_body("For each lead discovered, the pipeline automatically gathers the following data:")

add_styled_table(
    ["Data Point", "Source", "Purpose"],
    [
        ["Verified email", "Hunter.io, Apollo, Clearbit", "Email outreach"],
        ["Phone number", "Apollo, ZoomInfo", "Cold calling"],
        ["Company revenue & headcount", "Clearbit, Apollo", "Qualification"],
        ["Tech stack", "BuiltWith, Wappalyzer, job postings", "Personalization"],
        ["Recent funding", "Crunchbase API", "Budget signal"],
        ["Recent job postings", "LinkedIn, Indeed", "Need signal"],
        ["Company news", "Google News API", "Conversation starter"],
        ["LinkedIn activity", "LinkedIn", "Engagement hooks"],
        ["Mutual connections", "LinkedIn", "Warm intro paths"],
        ["Existing vendors", "G2, case studies", "Competitive positioning"],
    ],
    col_widths=[4, 5, 4],
)

add_heading_styled("4.2 AI-Powered Lead Scoring", 2)

add_body(
    "Each lead is scored 1-100 using the Claude API. The AI evaluates five dimensions "
    "and assigns a composite score that determines the outreach strategy:"
)

add_styled_table(
    ["Scoring Dimension", "Weight", "What It Measures"],
    [
        ["Decision-making authority", "0-25 points", "Title seniority and purchasing influence"],
        ["Industry fit", "0-20 points", "Alignment with AnalyticsGear's target verticals"],
        ["Tech stack alignment", "0-20 points", "Use of Snowflake, Databricks, dbt, Airflow, cloud platforms"],
        ["Intent signals", "0-20 points", "Hiring activity, funding, tech adoption, job postings"],
        ["Company size fit", "0-15 points", "100-5,000 employee sweet spot"],
    ],
    col_widths=[4.5, 3, 9],
)

add_heading_styled("4.3 Lead Tiers & Actions", 2)

add_styled_table(
    ["Tier", "Score Range", "Expected Volume", "Action"],
    [
        ["HOT", "75 - 100", "5-10 / day", "Personalized email + LinkedIn + phone call within 24h"],
        ["WARM", "50 - 74", "15-25 / day", "Personalized email + LinkedIn connect within 48h"],
        ["COLD", "25 - 49", "10-20 / day", "Template email + LinkedIn connect within 1 week"],
        ["DISCARD", "0 - 24", "Variable", "Archive, do not contact"],
    ],
    col_widths=[2.5, 3, 3, 8],
)

add_page_break()

# ════════════════════════════════════════════════════════════════
#  5. PHASE 3 — OUTREACH AUTOMATION
# ════════════════════════════════════════════════════════════════

add_heading_styled("5. Phase 3 — Outreach Automation", 1)

add_body(
    "Outreach is executed across three channels simultaneously: email, LinkedIn, and phone. "
    "Each channel has its own sequence, and the AI generates personalized content for every touchpoint."
)

add_heading_styled("5.1 Email Outreach", 2)

add_heading_styled("Deliverability Setup (Critical)", 3)
add_bullet("Set up SPF, DKIM, DMARC records for the outreach domain")
add_bullet("Use a separate domain for cold outreach (e.g., outreach.analyticsgear.com)")
add_bullet("Warm up the sending domain for 2-3 weeks before full volume")
add_bullet("Max 30-50 cold emails per day per email account")
add_bullet("Rotate across 3-5 email accounts")
add_bullet("Track bounce rates — keep below 3%")

add_heading_styled("Email Sequence — HOT Leads (5 emails over 14 days)", 3)

add_styled_table(
    ["Day", "Email Type", "Purpose"],
    [
        ["Day 1", "Personalized intro referencing their specific challenge", "Open the door"],
        ["Day 3", "Share a relevant AnalyticsGear blog post", "Add value, build credibility"],
        ["Day 6", "Brief case-study style proof point", "Social proof"],
        ["Day 10", '"Did this land?" — short bump email', "Re-engage"],
        ["Day 14", "Breakup email — closing the loop", "Final touch, leave door open"],
    ],
    col_widths=[2, 9, 5],
)

add_heading_styled("Email Sequence — WARM Leads (3 emails over 10 days)", 3)

add_styled_table(
    ["Day", "Email Type", "Purpose"],
    [
        ["Day 1", "Personalized intro", "Open the door"],
        ["Day 5", "Share relevant insight or blog post", "Add value"],
        ["Day 10", "Breakup email", "Final touch"],
    ],
    col_widths=[2, 9, 5],
)

add_heading_styled("AI Email Generation", 3)
add_body(
    "Every email is generated by the Claude API using the lead's enrichment data. The AI is prompted to: "
    "keep emails under 120 words, avoid salesy language, reference something specific about the prospect's "
    "company, end with a soft CTA, and write in a peer-to-peer consultative tone."
)

add_heading_styled("5.2 LinkedIn Outreach", 2)

add_styled_table(
    ["Step", "Day", "Action"],
    [
        ["1", "Day 0", "View their profile (triggers 'who viewed' notification)"],
        ["2", "Day 1", "Send connection request with personalized note (under 300 chars)"],
        ["3", "Day 2", "Like/comment on their recent post (if any)"],
        ["4", "Day 3", "Once connected — send intro message"],
        ["5", "Day 7", "Share a relevant article or insight"],
        ["6", "Day 14", "Follow up if no response"],
    ],
    col_widths=[1.5, 2, 13],
)

add_body(
    "Tools: Phantombuster or Dripify for automation, or manual sending with AI-drafted messages. "
    "Daily limit: 20-25 connection requests to avoid account restrictions."
)

add_heading_styled("5.3 Cold Calling (Semi-Automated)", 2)

add_body(
    "Cold calling cannot be fully automated but is heavily AI-assisted. The pipeline generates "
    "a call brief for each lead containing:"
)

add_numbered([
    "Opening hook (10 seconds, referencing something relevant to them)",
    "Value proposition tailored to their industry and tech stack",
    "Three discovery questions",
    "Common objections with suggested responses",
    "CTA: suggest a 20-minute discovery call",
])

add_body(
    "Calls are scheduled in daily blocks (e.g., 10 AM - 12 PM) using JustCall or Twilio. "
    "Recordings are auto-transcribed and summarized by AI for CRM logging."
)

add_page_break()

# ════════════════════════════════════════════════════════════════
#  6. PHASE 4 — FOLLOW-UP & NURTURING
# ════════════════════════════════════════════════════════════════

add_heading_styled("6. Phase 4 — Follow-Up & Nurturing", 1)

add_heading_styled("6.1 Engagement Signal Tracking", 2)

add_styled_table(
    ["Signal", "Source", "Follow-Up Action"],
    [
        ["Email opened", "SendGrid / Instantly", "Wait for click or reply"],
        ["Email link clicked", "SendGrid / Instantly", "Send follow-up within 2 hours"],
        ["Email replied", "Inbox monitoring", "Alert founder, respond within 1 hour"],
        ["LinkedIn connection accepted", "LinkedIn", "Send intro DM"],
        ["LinkedIn message read", "LinkedIn", "Follow up after 3 days if no reply"],
        ["Visited website", "Google Analytics / Clearbit Reveal", "Send contextual email"],
        ["No response after full sequence", "CRM", "Move to monthly nurture list"],
    ],
    col_widths=[4, 4.5, 8],
)

add_heading_styled("6.2 AI-Powered Reply Handling", 2)

add_body("When a prospect replies, the AI classifies the response and drafts an appropriate reply:")

add_styled_table(
    ["Classification", "AI Action"],
    [
        ["INTERESTED", "Draft reply proposing specific times for a 20-min discovery call"],
        ["OBJECTION", "Draft reply addressing the specific concern (budget, timing, vendor)"],
        ["NOT NOW", "Draft reply asking when to follow up, add to nurture list"],
        ["NOT INTERESTED", "Archive, send graceful close"],
        ["WRONG PERSON", "Draft reply asking for the right contact"],
        ["OUT OF OFFICE", "Reschedule follow-up for return date"],
    ],
    col_widths=[4, 13],
)

add_heading_styled("6.3 Long-Term Nurture Campaigns", 2)

add_body("For leads that don't convert immediately, run a monthly nurture program:")
add_bullet("Monthly insight email sharing new AnalyticsGear blog posts or case studies")
add_bullet('Quarterly check-in: "Anything changed on the data front?"')
add_bullet("Event invitations: webinars, AMAs, community events")
add_bullet("Re-engagement triggers: job changes, new funding, posts about data challenges")

add_page_break()

# ════════════════════════════════════════════════════════════════
#  7. PHASE 5 — CRM
# ════════════════════════════════════════════════════════════════

add_heading_styled("7. Phase 5 — CRM, Logging & Tracking", 1)

add_heading_styled("7.1 CRM Platform Recommendation", 2)

add_styled_table(
    ["Option", "Cost", "Best For"],
    [
        ["Google Sheets + Apps Script", "Free", "Starting out, maximum flexibility"],
        ["Notion Database", "Free - $10/mo", "Visual pipeline, team collaboration"],
        ["HubSpot CRM (Free tier)", "Free", "Proper CRM with email integration"],
        ["Airtable", "Free - $20/mo", "Flexible database with automations"],
    ],
    col_widths=[5, 3.5, 8],
)

add_callout_box(
    "Start with Google Sheets for speed and flexibility. Migrate to HubSpot Free CRM "
    "once lead volume exceeds 200 active leads.",
    title="Recommendation:"
)

add_heading_styled("7.2 CRM Schema — Master Lead List", 2)

add_styled_table(
    ["Field", "Description"],
    [
        ["Lead ID", "Auto-generated unique identifier"],
        ["Name / Email / Phone / LinkedIn", "Core contact information"],
        ["Company / Title / Industry", "Firmographic data"],
        ["Company Size", "Employee headcount"],
        ["Tech Stack", "Detected technologies"],
        ["Lead Source", "LinkedIn / Apollo / Job Board / Inbound / Event / Social"],
        ["Lead Score (1-100)", "AI-generated composite score"],
        ["Lead Tier", "HOT / WARM / COLD"],
        ["Pain Signals", "Detected challenges and needs"],
        ["Status", "New / Contacted / Replied / Meeting / Proposal / Won / Lost"],
        ["Current Sequence & Step", "Which outreach sequence, which step"],
        ["Date Added / Last Contacted", "Timeline tracking"],
        ["Next Follow-Up", "Scheduled next action date"],
        ["Notes", "AI-generated + manual notes"],
    ],
    col_widths=[5, 12],
)

add_heading_styled("7.3 Activity Log Schema", 2)

add_styled_table(
    ["Field", "Description"],
    [
        ["Timestamp", "When the activity occurred"],
        ["Lead ID", "Link to master lead record"],
        ["Channel", "Email / LinkedIn / Phone / Website"],
        ["Action", "Sent / Opened / Clicked / Replied / Called / Connected"],
        ["Details", "Email subject, message content, call notes"],
        ["Outcome", "Positive / Neutral / Negative / No Response"],
    ],
    col_widths=[4, 13],
)

add_page_break()

# ════════════════════════════════════════════════════════════════
#  8. PHASE 6 — ANALYTICS
# ════════════════════════════════════════════════════════════════

add_heading_styled("8. Phase 6 — Analytics & Optimization", 1)

add_heading_styled("8.1 Key Performance Metrics", 2)

add_styled_table(
    ["Metric", "Target", "Frequency"],
    [
        ["Leads discovered per week", "50 - 100", "Weekly"],
        ["Emails sent per day", "30 - 50", "Daily"],
        ["Email open rate", "> 40%", "Weekly"],
        ["Email reply rate", "> 5%", "Weekly"],
        ["LinkedIn connection acceptance rate", "> 30%", "Weekly"],
        ["LinkedIn reply rate", "> 10%", "Weekly"],
        ["Meetings booked per week", "3 - 5", "Weekly"],
        ["Meeting to proposal rate", "> 40%", "Monthly"],
        ["Proposal to close rate", "> 25%", "Monthly"],
        ["Cost per meeting booked", "< $50", "Monthly"],
        ["Pipeline value", "Growing MoM", "Monthly"],
    ],
    col_widths=[5.5, 3, 3],
)

add_heading_styled("8.2 Weekly AI Analysis Report", 2)

add_body(
    "Every Friday, the pipeline generates an AI-powered weekly report that includes:"
)

add_numbered([
    "Executive summary (3 sentences on pipeline health)",
    "What's working: top performing channels, messages, and industries",
    "What's not working: low engagement areas requiring attention",
    "Specific recommendations for the following week",
    "A/B test suggestions for email subject lines and messaging",
])

add_body("The report is automatically sent to the founders via email or Slack.")

add_page_break()

# ════════════════════════════════════════════════════════════════
#  9. TECH STACK
# ════════════════════════════════════════════════════════════════

add_heading_styled("9. Tech Stack & Tools", 1)

add_styled_table(
    ["Category", "Tool", "Purpose", "Cost/Month"],
    [
        ["Orchestration", "Python + Cron", "Run daily pipeline scripts", "Free"],
        ["Orchestration", "n8n (self-hosted) or Prefect", "Visual workflow, error handling", "Free"],
        ["Prospecting", "Apollo.io (Basic)", "Lead database + emails + enrichment", "$49"],
        ["Prospecting", "Apify (Starter)", "Web scraping (LinkedIn, job boards)", "$49"],
        ["Prospecting", "LinkedIn Sales Navigator", "Advanced lead search", "$80"],
        ["Email", "Instantly.ai", "Cold email sending + warmup", "$30"],
        ["Email", "Mailgun / SendGrid", "Transactional email + tracking", "Free tier"],
        ["LinkedIn", "Phantombuster", "LinkedIn automation", "$56"],
        ["AI Engine", "Claude API (Anthropic)", "Personalization, scoring, analysis", "$30-50"],
        ["CRM", "Google Sheets", "Lead tracking and logging", "Free"],
        ["Phone", "JustCall / Twilio", "Call scheduling, auto-dial", "$19"],
        ["Monitoring", "Google Analytics", "Website visitor tracking", "Free"],
        ["Domain", "Cold outreach domain", "Protect main domain reputation", "$1"],
    ],
    col_widths=[3, 4.5, 5.5, 3],
)

add_callout_box(
    "Estimated total: $314 - $334 per month\n"
    "Compare to: $3,000 - $5,000 per month for a junior SDR hire\n"
    "Savings: approximately 90%",
    title="Total Monthly Cost:"
)

add_page_break()

# ════════════════════════════════════════════════════════════════
#  10. DAILY SCHEDULE
# ════════════════════════════════════════════════════════════════

add_heading_styled("10. Daily Automation Schedule", 1)

add_body("The pipeline runs automatically each day on the following schedule:")

add_styled_table(
    ["Time", "Phase", "Tasks"],
    [
        ["06:00 AM", "Prospecting", "Run Apollo.io search, LinkedIn Navigator export, job board scraping, deduplicate against CRM"],
        ["07:00 AM", "Enrichment", "Verify emails (Hunter.io), gather company data (Clearbit), AI lead scoring, assign tiers"],
        ["08:00 AM", "Outreach", "Generate personalized emails, send via Instantly.ai, LinkedIn connection requests, generate call briefs"],
        ["09:00 AM", "Follow-Up", "Check email opens/clicks/replies, LinkedIn acceptances, trigger follow-up sequences, AI-classify replies"],
        ["09:30 AM", "Logging", "Log all activities to Google Sheets, update lead statuses, calculate next follow-up dates"],
        ["10:00 AM", "Human Review", "Review AI-drafted replies (15-30 mins), make cold calls from briefs, approve HOT lead outreach"],
        ["06:00 PM", "Daily Summary", "AI generates end-of-day report: leads discovered, contacted, replied, meetings booked, tomorrow's priorities"],
    ],
    col_widths=[2.5, 3, 11],
)

add_callout_box(
    "The founder's daily time commitment is approximately 30-60 minutes, "
    "primarily for reviewing AI-drafted replies, making phone calls, and "
    "approving outreach to high-value leads. Everything else runs on autopilot.",
    title="Founder Time Required:"
)

add_page_break()

# ════════════════════════════════════════════════════════════════
#  11. IMPLEMENTATION ROADMAP
# ════════════════════════════════════════════════════════════════

add_heading_styled("11. Implementation Roadmap", 1)

add_heading_styled("Weeks 1-2: Foundation", 2)
checklist_1 = [
    "Set up separate cold outreach domain (e.g., outreach.analyticsgear.com)",
    "Configure SPF, DKIM, DMARC for email deliverability",
    "Create Apollo.io account, configure ICP filters",
    "Set up Instantly.ai, begin email warmup (takes 2-3 weeks)",
    "Create Google Sheets CRM with the schema from Section 7",
    "Set up Claude API key, test lead scoring prompts",
    "Write core Python scripts: prospecting, enrichment, scoring",
    "Define ICP and build initial target company list (200 companies)",
]
for item in checklist_1:
    add_bullet(item)

add_heading_styled("Weeks 3-4: Build Outreach Engine", 2)
checklist_2 = [
    "Write email generation prompts, test with 10 sample leads",
    "Set up Phantombuster for LinkedIn automation",
    "Build email sequence templates (HOT, WARM, COLD)",
    "Build LinkedIn sequence automation",
    "Write activity logging scripts to Google Sheets",
    "Create daily summary report generator",
    "Set up cron jobs for daily pipeline execution",
    "Test full pipeline end-to-end with 20 real leads",
]
for item in checklist_2:
    add_bullet(item)

add_heading_styled("Weeks 5-6: Launch & Iterate", 2)
checklist_3 = [
    "Begin sending cold emails (start at 10/day, ramp to 50/day)",
    "Begin LinkedIn outreach (10-20 connections/day)",
    "Monitor deliverability, open rates, reply rates",
    "Set up cold calling workflow with JustCall",
    "Start making 5-10 calls/day using AI call briefs",
    "Review and respond to replies daily",
    "Iterate on email copy based on performance data",
]
for item in checklist_3:
    add_bullet(item)

add_heading_styled("Weeks 7-8: Optimize & Scale", 2)
checklist_4 = [
    "Analyze first month's data: which channels and messages work best",
    "A/B test subject lines, email copy, LinkedIn messages",
    "Refine ICP based on actual engagement data",
    "Add new lead sources (events, social listening, intent data)",
    "Build monthly nurture campaign for non-responders",
    "Consider migrating to HubSpot Free CRM if volume grows",
    "Document SOPs for any manual steps",
]
for item in checklist_4:
    add_bullet(item)

add_heading_styled("Ongoing (Month 3+)", 2)
add_bullet("Weekly: Review pipeline metrics, adjust targeting")
add_bullet("Monthly: AI-generated performance report + strategy recommendations")
add_bullet("Quarterly: Full pipeline audit, add new channels or sources")
add_bullet("Continuous: Improve AI prompts based on conversion data")

add_page_break()

# ════════════════════════════════════════════════════════════════
#  12. COST ESTIMATES
# ════════════════════════════════════════════════════════════════

add_heading_styled("12. Cost Estimates", 1)

add_heading_styled("12.1 Monthly Operating Costs", 2)

add_styled_table(
    ["Item", "Monthly Cost", "Notes"],
    [
        ["Apollo.io (Basic)", "$49", "Lead database + email finder"],
        ["Instantly.ai", "$30", "Cold email platform + warmup"],
        ["LinkedIn Sales Navigator", "$80", "Advanced search"],
        ["Phantombuster", "$56", "LinkedIn automation"],
        ["Claude API (Anthropic)", "$30 - $50", "~100K tokens/day for personalization"],
        ["JustCall (Essentials)", "$19", "Cold calling"],
        ["Apify (Starter)", "$49", "Web scraping"],
        ["Cold outreach domain", "$1", "~$12/year"],
        ["TOTAL", "$314 - $334", ""],
    ],
    col_widths=[5, 3.5, 8],
)

add_heading_styled("12.2 AI Pipeline vs. Sales Hire Comparison", 2)

add_styled_table(
    ["Factor", "AI Pipeline", "Junior SDR Hire"],
    [
        ["Monthly cost", "~$350", "$3,000 - $5,000"],
        ["Works on schedule 24/7", "Yes", "No"],
        ["Scales instantly", "Add more accounts", "Hire more people"],
        ["Consistent quality", "Yes", "Variable"],
        ["Personal touch", "Lower (needs human review)", "Higher"],
        ["Relationship building", "Limited", "Strong"],
        ["Annual cost", "~$4,200", "$36,000 - $60,000"],
    ],
    col_widths=[4.5, 5.5, 5.5],
)

add_callout_box(
    "The AI pipeline costs approximately 90% less and handles 80% of a sales person's tasks. "
    "The remaining 20% (calls, relationship building, closing) requires 30-60 minutes of "
    "founder time per day.",
    title="Bottom Line:"
)

add_page_break()

# ════════════════════════════════════════════════════════════════
#  13. COMPLIANCE
# ════════════════════════════════════════════════════════════════

add_heading_styled("13. Risk Mitigation & Compliance", 1)

add_heading_styled("13.1 Email Compliance", 2)
add_bullet("CAN-SPAM (US): ", "Include unsubscribe link, physical address, no deceptive subject lines")
add_bullet("GDPR (EU): ", "Legitimate interest basis for B2B outreach, honor opt-outs within 30 days")
add_bullet("CASL (Canada): ", "Requires implied or express consent — be cautious with Canadian leads")
add_bullet("Best practice: ", "Include easy opt-out in every email, honor requests immediately")

add_heading_styled("13.2 LinkedIn Compliance", 2)
add_bullet("Daily limit: 20-25 connection requests per day to avoid restrictions")
add_bullet("Use human-like delays (random intervals) to avoid automation detection")
add_bullet("Keep LinkedIn profile professional and complete")
add_bullet("Fallback: if automation gets restricted, switch to manual + AI-assisted approach")

add_heading_styled("13.3 Domain & Reputation Protection", 2)
add_bullet("Use a separate domain for cold outreach to protect analyticsgear.com")
add_bullet("Monitor bounce rates — remove invalid emails immediately")
add_bullet("Warm up email accounts for 2-3 weeks before sending at volume")
add_bullet("Check blacklists weekly using MXToolbox and Spamhaus")

add_heading_styled("13.4 Data Privacy", 2)
add_bullet("Don't store sensitive personal data beyond what's needed for outreach")
add_bullet("Regular cleanup of old/unresponsive leads (90-day retention for non-engaged)")
add_bullet("Secure access to CRM data (limited to team members only)")

add_page_break()

# ════════════════════════════════════════════════════════════════
#  14. TEMPLATES
# ════════════════════════════════════════════════════════════════

add_heading_styled("14. Email & LinkedIn Templates", 1)

add_heading_styled("14.1 Cold Email Templates", 2)

add_heading_styled("Template 1: Pain-Point Hook (Data Engineering)", 3)

add_callout_box(
    'Subject: {company}\'s data stack\n\n'
    'Hi {first_name},\n\n'
    'Noticed {company} is hiring for data engineering roles — usually means the pipeline '
    'backlog is growing faster than the team.\n\n'
    'We\'ve helped {similar_company} clear a 6-month backlog in 8 weeks by embedding senior '
    'data engineers alongside their team (Snowflake + dbt + Airflow).\n\n'
    'If the data engineering bottleneck is real, happy to share what worked for them — '
    '15 minutes, no pitch.\n\n'
    'Best,\n{sender_name}\nAnalyticsGear'
)

add_heading_styled("Template 2: Tech-Stack Specific", 3)

add_callout_box(
    'Subject: quick question about {tech_platform} at {company}\n\n'
    'Hi {first_name},\n\n'
    'Saw that {company} is running on {tech_platform}. We specialize in optimizing '
    '{tech_platform} deployments — things like query cost reduction, pipeline reliability, '
    'and proper data modeling.\n\n'
    'One client cut their {tech_platform} spend by 40% after we restructured their '
    'warehouse — took 4 weeks.\n\n'
    'Worth a quick chat to see if there\'s a similar opportunity?\n\n'
    '{sender_name}'
)

add_heading_styled("Template 3: Industry-Specific (Banking/Finance)", 3)

add_callout_box(
    'Subject: data compliance at {company}\n\n'
    'Hi {first_name},\n\n'
    'Data teams in banking are dealing with a unique headache right now — regulatory pressure '
    'to implement data lineage and quality frameworks, while also shipping analytics faster.\n\n'
    'We\'ve worked with financial services firms to build compliant data platforms that don\'t '
    'slow teams down. Happy to share the playbook if that resonates.\n\n'
    '{sender_name}\nAnalyticsGear'
)

add_heading_styled("14.2 LinkedIn Message Templates", 2)

add_heading_styled("Connection Accepted — Intro Message", 3)

add_callout_box(
    'Thanks for connecting, {first_name}!\n\n'
    'I work with {industry} companies on their data and AI platforms. Your team\'s work '
    'at {company} caught my eye — particularly {specific_detail}.\n\n'
    'Would love to hear how you\'re thinking about {relevant_topic} on your end. '
    'No pitch, genuinely curious.'
)

add_heading_styled("Follow-Up After No Response", 3)

add_callout_box(
    'Hi {first_name} — don\'t want to be that person who sends five follow-ups. '
    'Just wanted to share this quick read our team published on {relevant_topic}:\n\n'
    '{blog_link}\n\n'
    'Thought it might be relevant given what {company} is building. '
    'Either way, great to be connected.'
)

add_page_break()

# ════════════════════════════════════════════════════════════════
#  QUICK START CHECKLIST (Final page)
# ════════════════════════════════════════════════════════════════

add_heading_styled("Quick Start Checklist", 1)

add_body("Follow these 10 steps to get the pipeline operational:")

steps = [
    ("Today", "Sign up for Apollo.io (free trial), Instantly.ai, and Claude API"),
    ("This Week", "Buy cold outreach domain, set up DNS records (SPF/DKIM/DMARC), start email warmup"),
    ("This Week", "Build Google Sheets CRM using the schema from Section 7"),
    ("Week 2", "Write first Python scripts (prospecting + scoring)"),
    ("Week 2", "Generate first batch of 50 leads, score them, review quality"),
    ("Week 3", "Send first 10 cold emails (manually review AI-generated copy)"),
    ("Week 3", "Begin LinkedIn outreach (10 connections/day)"),
    ("Week 4", "Automate with cron, begin daily pipeline runs"),
    ("Month 2", "Optimize based on data — identify what's generating replies"),
    ("Month 3", "Full autopilot with 30-minute daily founder review"),
]

for i, (when, what) in enumerate(steps, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"  {i:2d}. ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = INDIGO
    r2 = p.add_run(f"[{when}]  ")
    r2.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = TEAL
    r3 = p.add_run(what)
    r3.font.size = Pt(10)
    r3.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(6)

doc.add_paragraph()
doc.add_paragraph()

# Footer note
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("━" * 40)
r.font.color.rgb = INDIGO
r.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "This blueprint was designed for AnalyticsGear's specific services, ICP, and growth stage.\n"
    "The pipeline should be treated as a living system — continuously improved based on real engagement data."
)
r.font.size = Pt(9)
r.font.color.rgb = GRAY
r.italic = True

# ── Save ──
output_path = r"c:\analyticsgear\sales_pipeline\AnalyticsGear_AI_Sales_Pipeline_Blueprint.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
